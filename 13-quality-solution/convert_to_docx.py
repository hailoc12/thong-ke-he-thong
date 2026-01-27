#!/usr/bin/env python3
"""
Convert Strategic Dashboard Feature Description to DOCX
Format: Nghị định 30/2020/NĐ-CP
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


def setup_document_format(doc):
    """Setup document format theo Nghị định 30"""
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # Setup default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_before = Pt(6)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Setup heading styles
    for i in range(1, 5):
        try:
            heading_style = doc.styles[f'Heading {i}']
        except KeyError:
            heading_style = doc.styles.add_style(f'Heading {i}', WD_STYLE_TYPE.PARAGRAPH)

        heading_style.font.name = 'Times New Roman'
        heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        heading_style.font.bold = True
        heading_style.font.underline = False
        heading_style.paragraph_format.line_spacing = 1.5
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)

        if i == 1:
            heading_style.font.size = Pt(16)
        elif i == 2:
            heading_style.font.size = Pt(15)
        elif i == 3:
            heading_style.font.size = Pt(14)
        else:
            heading_style.font.size = Pt(14)


def parse_inline_markdown(text, paragraph):
    """Parse inline markdown: **bold**, *italic*, `code`"""
    # Pattern to find markdown elements
    patterns = [
        (r'\*\*\*(.+?)\*\*\*', 'bold_italic'),  # ***bold italic***
        (r'\*\*(.+?)\*\*', 'bold'),              # **bold**
        (r'\*(.+?)\*', 'italic'),                # *italic*
        (r'`(.+?)`', 'code'),                    # `code`
    ]

    # Simple approach: handle bold first
    remaining_text = text

    # Split by bold markers
    bold_pattern = r'\*\*(.+?)\*\*'
    parts = re.split(bold_pattern, remaining_text)

    is_bold = False
    for i, part in enumerate(parts):
        if not part:
            continue

        if i % 2 == 1:  # This is bold text
            run = paragraph.add_run(part)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
        else:  # Normal text - check for code
            code_pattern = r'`(.+?)`'
            code_parts = re.split(code_pattern, part)

            for j, code_part in enumerate(code_parts):
                if not code_part:
                    continue

                if j % 2 == 1:  # This is code
                    run = paragraph.add_run(code_part)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)  # Dark red
                else:  # Normal text
                    run = paragraph.add_run(code_part)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)


def add_table_from_markdown(doc, table_lines):
    """Convert markdown table to DOCX table"""
    # Parse header
    header_line = table_lines[0]
    headers = [h.strip() for h in header_line.split('|')[1:-1]]

    # Parse data rows (skip separator line at index 1)
    data_rows = []
    for line in table_lines[2:]:
        if line.strip():
            cells = [c.strip() for c in line.split('|')[1:-1]]
            data_rows.append(cells)

    # Create table
    num_rows = len(data_rows) + 1
    num_cols = len(headers)

    if num_cols == 0:
        return

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.paragraphs[0].clear()

        # Add shading to header
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'E3F2FD')
        cell._tc.get_or_add_tcPr().append(shading_elm)

        # Clean header text (remove markdown)
        clean_header = re.sub(r'\*\*(.+?)\*\*', r'\1', header)
        run = cell.paragraphs[0].add_run(clean_header)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for row_idx, row_data in enumerate(data_rows):
        cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < len(cells):
                cell = cells[col_idx]
                cell.paragraphs[0].clear()

                # Parse inline markdown in cell
                clean_text = cell_text.strip()
                parse_inline_markdown(clean_text, cell.paragraphs[0])

                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Set column widths
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(1.5)

    return table


def add_code_block(doc, code_lines):
    """Add code block as formatted paragraph"""
    code_text = '\n'.join(code_lines)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # Add border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), 'CCCCCC')
        pBdr.append(border)
    pPr.append(pBdr)

    # Add shading
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shading)

    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(11)


def convert_markdown_to_docx(md_content, output_path):
    """Main conversion function"""
    doc = Document()
    setup_document_format(doc)

    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]

        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                add_code_block(doc, code_block_lines)
                code_block_lines = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        # Handle tables
        if line.strip().startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)

            # Check if next line is also table
            if i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
                i += 1
                continue
            else:
                # End of table
                if len(table_lines) >= 2:
                    add_table_from_markdown(doc, table_lines)
                    doc.add_paragraph()  # Add spacing after table
                table_lines = []
                in_table = False
                i += 1
                continue

        # Handle horizontal rule
        if line.strip() == '---':
            # Add subtle separator
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue

        # Handle headings
        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_paragraph(style='Heading 1')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.bold = True
            i += 1
            continue

        if line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_paragraph(style='Heading 2')
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.bold = True
            i += 1
            continue

        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_paragraph(style='Heading 3')
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.bold = True
            i += 1
            continue

        if line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_paragraph(style='Heading 4')
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.bold = True
            i += 1
            continue

        # Handle bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.first_line_indent = Cm(-0.5)

            # Add bullet
            bullet_run = p.add_run('• ')
            bullet_run.font.name = 'Times New Roman'
            bullet_run.font.size = Pt(14)

            parse_inline_markdown(text, p)
            i += 1
            continue

        # Handle numbered lists
        num_match = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if num_match:
            number = num_match.group(1)
            text = num_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.first_line_indent = Cm(-0.5)

            # Add number
            num_run = p.add_run(f'{number}. ')
            num_run.font.name = 'Times New Roman'
            num_run.font.size = Pt(14)

            parse_inline_markdown(text, p)
            i += 1
            continue

        # Handle blockquote (italic centered text at the end)
        if line.strip().startswith('*"') and line.strip().endswith('"*'):
            text = line.strip()[2:-2]  # Remove *" and "*
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(18)
            run = p.add_run(text)
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            i += 1
            continue

        # Handle empty lines
        if not line.strip():
            i += 1
            continue

        # Handle normal paragraphs
        p = doc.add_paragraph()
        parse_inline_markdown(line, p)

        i += 1

    # Save document
    doc.save(output_path)
    print(f"✓ Đã tạo file: {output_path}")


def main():
    # Input/output paths
    input_path = Path(__file__).parent / 'strategic-dashboard-feature-description.md'
    output_path = Path(__file__).parent / 'strategic-dashboard-feature-description.docx'

    # Read markdown content
    with open(input_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Convert
    convert_markdown_to_docx(md_content, str(output_path))

    print(f"\n✓ Chuyển đổi hoàn tất!")
    print(f"  Input:  {input_path}")
    print(f"  Output: {output_path}")


if __name__ == '__main__':
    main()
