#!/usr/bin/env python3
"""
Convert Strategic Dashboard Feature Description v2 to Professional DOCX
- Renders Mermaid diagrams to PNG images
- Formats according to Nghị định 30/2020/NĐ-CP
- Processes markdown tables and inline formatting
"""

import re
import os
import subprocess
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def setup_document_format(doc):
    """Setup document format theo Nghị định 30"""
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # Setup default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_before = Pt(6)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Setup heading styles
    for i in range(1, 5):
        try:
            heading_style = doc.styles[f'Heading {i}']
        except KeyError:
            heading_style = doc.styles.add_style(f'Heading {i}', WD_STYLE_TYPE.PARAGRAPH)

        heading_style.font.name = 'Times New Roman'
        heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        heading_style.font.bold = True
        heading_style.font.underline = False
        heading_style.paragraph_format.line_spacing = 1.5
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)

        if i == 1:
            heading_style.font.size = Pt(16)
        elif i == 2:
            heading_style.font.size = Pt(15)
        elif i == 3:
            heading_style.font.size = Pt(14)
        else:
            heading_style.font.size = Pt(14)


def render_mermaid_diagram(mermaid_code, output_path, diagram_index):
    """Render Mermaid diagram to PNG using mmdc CLI"""
    mmd_file = output_path / f'diagram_{diagram_index}.mmd'
    png_file = output_path / f'diagram_{diagram_index}.png'

    # Write mermaid code to file
    with open(mmd_file, 'w', encoding='utf-8') as f:
        f.write(mermaid_code)

    try:
        # Try to render using npx mmdc
        result = subprocess.run(
            ['npx', '-p', '@mermaid-js/mermaid-cli', 'mmdc',
             '-i', str(mmd_file),
             '-o', str(png_file),
             '-b', 'white',
             '-w', '1200'],
            capture_output=True,
            text=True,
            timeout=60
        )

        if result.returncode == 0 and png_file.exists():
            print(f"✓ Rendered diagram {diagram_index}: {png_file}")
            return str(png_file)
        else:
            print(f"✗ Failed to render diagram {diagram_index}: {result.stderr}")
            return None
    except subprocess.TimeoutExpired:
        print(f"✗ Timeout rendering diagram {diagram_index}")
        return None
    except FileNotFoundError:
        print(f"✗ mmdc not found, trying alternative method")
        return None


def extract_mermaid_diagrams(md_content):
    """Extract Mermaid diagrams from markdown content"""
    pattern = r'```mermaid\n(.*?)```'
    matches = re.findall(pattern, md_content, re.DOTALL)
    return matches


def parse_inline_markdown(text, paragraph):
    """Parse inline markdown: **bold**, *italic*, `code`"""
    # Split by bold markers
    bold_pattern = r'\*\*(.+?)\*\*'
    parts = re.split(bold_pattern, text)

    for i, part in enumerate(parts):
        if not part:
            continue

        if i % 2 == 1:  # This is bold text
            run = paragraph.add_run(part)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
        else:  # Normal text - check for code
            code_pattern = r'`(.+?)`'
            code_parts = re.split(code_pattern, part)

            for j, code_part in enumerate(code_parts):
                if not code_part:
                    continue

                if j % 2 == 1:  # This is code
                    run = paragraph.add_run(code_part)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
                else:  # Normal text
                    run = paragraph.add_run(code_part)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)


def add_table_from_markdown(doc, table_lines):
    """Convert markdown table to DOCX table"""
    header_line = table_lines[0]
    headers = [h.strip() for h in header_line.split('|')[1:-1]]

    data_rows = []
    for line in table_lines[2:]:
        if line.strip():
            cells = [c.strip() for c in line.split('|')[1:-1]]
            data_rows.append(cells)

    num_rows = len(data_rows) + 1
    num_cols = len(headers)

    if num_cols == 0:
        return

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.paragraphs[0].clear()

        # Add shading to header
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'E3F2FD')
        cell._tc.get_or_add_tcPr().append(shading_elm)

        clean_header = re.sub(r'\*\*(.+?)\*\*', r'\1', header)
        run = cell.paragraphs[0].add_run(clean_header)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for row_idx, row_data in enumerate(data_rows):
        cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < len(cells):
                cell = cells[col_idx]
                cell.paragraphs[0].clear()
                parse_inline_markdown(cell_text.strip(), cell.paragraphs[0])
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    return table


def add_diagram_placeholder(doc, diagram_index, diagram_code):
    """Add a text placeholder for diagram that couldn't be rendered"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    run = p.add_run(f"[Hình {diagram_index}: Sơ đồ - xem file gốc để biết chi tiết]")
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def convert_markdown_to_docx(md_content, output_path, diagrams_dir):
    """Main conversion function"""
    doc = Document()
    setup_document_format(doc)

    # Extract and render diagrams first
    mermaid_diagrams = extract_mermaid_diagrams(md_content)
    diagram_images = {}

    print(f"\nFound {len(mermaid_diagrams)} Mermaid diagrams")

    for idx, diagram_code in enumerate(mermaid_diagrams, 1):
        png_path = render_mermaid_diagram(diagram_code, diagrams_dir, idx)
        if png_path:
            diagram_images[idx] = png_path

    # Replace mermaid blocks with placeholders
    diagram_idx = 0
    def replace_mermaid(match):
        nonlocal diagram_idx
        diagram_idx += 1
        return f'{{{{DIAGRAM_{diagram_idx}}}}}'

    content_with_placeholders = re.sub(r'```mermaid\n.*?```', replace_mermaid, md_content, flags=re.DOTALL)

    lines = content_with_placeholders.split('\n')
    i = 0
    in_table = False
    table_lines = []
    current_diagram = 1

    while i < len(lines):
        line = lines[i]

        # Handle diagram placeholders
        diagram_match = re.match(r'\{\{DIAGRAM_(\d+)\}\}', line.strip())
        if diagram_match:
            diagram_num = int(diagram_match.group(1))
            if diagram_num in diagram_images:
                # Add image
                doc.add_paragraph()  # spacing
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                try:
                    run.add_picture(diagram_images[diagram_num], width=Inches(5.5))
                except Exception as e:
                    print(f"Error adding image: {e}")
                    add_diagram_placeholder(doc, diagram_num, "")

                # Add caption
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption.add_run(f"Hình {diagram_num}")
                caption_run.italic = True
                caption_run.font.name = 'Times New Roman'
                caption_run.font.size = Pt(12)
            else:
                add_diagram_placeholder(doc, diagram_num, "")
            i += 1
            continue

        # Handle code blocks (non-mermaid)
        if line.strip().startswith('```') and 'mermaid' not in line:
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```

            # Add code block with formatting
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.line_spacing = 1.0

            # Add border and shading
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '4')
                border.set(qn('w:color'), 'CCCCCC')
                pBdr.append(border)
            pPr.append(pBdr)

            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F5F5F5')
            pPr.append(shading)

            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(11)
            continue

        # Handle tables
        if line.strip().startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)

            if i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
                i += 1
                continue
            else:
                if len(table_lines) >= 2:
                    add_table_from_markdown(doc, table_lines)
                    doc.add_paragraph()
                table_lines = []
                in_table = False
                i += 1
                continue

        # Handle horizontal rule
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue

        # Handle headings
        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_paragraph(style='Heading 1')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.bold = True
            i += 1
            continue

        if line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_paragraph(style='Heading 2')
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.bold = True
            i += 1
            continue

        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_paragraph(style='Heading 3')
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.bold = True
            i += 1
            continue

        if line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_paragraph(style='Heading 4')
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.bold = True
            i += 1
            continue

        # Handle bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.first_line_indent = Cm(-0.5)

            bullet_run = p.add_run('• ')
            bullet_run.font.name = 'Times New Roman'
            bullet_run.font.size = Pt(14)

            parse_inline_markdown(text, p)
            i += 1
            continue

        # Handle numbered lists
        num_match = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if num_match:
            number = num_match.group(1)
            text = num_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.first_line_indent = Cm(-0.5)

            num_run = p.add_run(f'{number}. ')
            num_run.font.name = 'Times New Roman'
            num_run.font.size = Pt(14)

            parse_inline_markdown(text, p)
            i += 1
            continue

        # Handle blockquote (italic centered text)
        if line.strip().startswith('*"') and line.strip().endswith('"*'):
            text = line.strip()[2:-2]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(18)
            run = p.add_run(text)
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            i += 1
            continue

        # Handle empty lines
        if not line.strip():
            i += 1
            continue

        # Handle normal paragraphs
        p = doc.add_paragraph()
        parse_inline_markdown(line, p)

        i += 1

    # Save document
    doc.save(output_path)
    print(f"\n✓ Đã tạo file: {output_path}")


def main():
    base_dir = Path(__file__).parent
    input_path = base_dir / 'strategic-dashboard-feature-v2.md'
    output_path = base_dir / 'strategic-dashboard-feature-v2.docx'
    diagrams_dir = base_dir / 'diagrams'

    # Create diagrams directory
    diagrams_dir.mkdir(exist_ok=True)

    # Read markdown content
    with open(input_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Convert
    convert_markdown_to_docx(md_content, str(output_path), diagrams_dir)

    print(f"\n✓ Chuyển đổi hoàn tất!")
    print(f"  Input:  {input_path}")
    print(f"  Output: {output_path}")
    print(f"  Diagrams: {diagrams_dir}")


if __name__ == '__main__':
    main()
