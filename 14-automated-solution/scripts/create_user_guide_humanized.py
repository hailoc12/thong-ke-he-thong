#!/usr/bin/env python3
"""
Script tạo Hướng dẫn sử dụng - Phiên bản Humanized
Tuân thủ Nghị định 30/2020/NĐ-CP về công tác văn thư
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import os

# Path to screenshots
SCREENSHOTS_DIR = ".playwright-mcp/screenshots"

def setup_document_format(doc):
    """Thiết lập format theo Nghị định 30/2020/NĐ-CP"""
    # Thiết lập margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # Thiết lập style Normal
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Paragraph format
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.space_before = Pt(6)
    paragraph_format.space_after = Pt(6)
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def configure_heading_styles(doc):
    """Cấu hình các style Heading theo chuẩn"""
    # Heading 1 - Tiêu đề chương
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 51, 102)  # Xanh đậm
    h1.font.underline = False  # KHÔNG gạch chân
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    h1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Heading 2 - Tiêu đề mục
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 51, 102)
    h2.font.underline = False
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(6)
    h2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Heading 3 - Tiêu đề nhỏ
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.italic = True
    h3.font.underline = False
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def add_image_with_caption(doc, image_path, caption, width=Inches(5.5)):
    """Thêm hình ảnh với chú thích"""
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Chú thích hình
        caption_para = doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption_para.add_run(caption)
        run.italic = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[Minh họa: {caption}]")
        run.italic = True

def add_note_box(doc, note_text, note_type="info"):
    """Thêm khung lưu ý nổi bật"""
    p = doc.add_paragraph()
    if note_type == "warning":
        run = p.add_run("⚠ Lưu ý quan trọng: ")
    elif note_type == "tip":
        run = p.add_run("💡 Mẹo: ")
    else:
        run = p.add_run("ℹ Ghi chú: ")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    text_run = p.add_run(note_text)
    text_run.font.name = 'Times New Roman'
    text_run.font.size = Pt(14)

def create_user_guide():
    """Tạo tài liệu hướng dẫn sử dụng"""
    doc = Document()
    setup_document_format(doc)
    configure_heading_styles(doc)

    # ==================== TRANG BÌA ====================
    doc.add_paragraph()
    doc.add_paragraph()

    # Quốc hiệu tiêu ngữ
    quoc_hieu = doc.add_paragraph()
    quoc_hieu.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = quoc_hieu.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM')
    run1.bold = True
    run1.font.size = Pt(13)
    run1.font.name = 'Times New Roman'

    tieu_ngu = doc.add_paragraph()
    tieu_ngu.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = tieu_ngu.add_run('Độc lập - Tự do - Hạnh phúc')
    run2.bold = True
    run2.font.size = Pt(13)
    run2.font.name = 'Times New Roman'

    # Gạch ngang
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = line.add_run('─────────────────')
    run_line.font.name = 'Times New Roman'

    doc.add_paragraph()
    doc.add_paragraph()

    # Tiêu đề chính
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run('HƯỚNG DẪN SỬ DỤNG')
    run_title.bold = True
    run_title.font.size = Pt(20)
    run_title.font.name = 'Times New Roman'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle.add_run('Hệ thống Khảo sát Chuyển đổi số trực tuyến')
    run_sub.bold = True
    run_sub.font.size = Pt(16)
    run_sub.font.name = 'Times New Roman'

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Thông tin đơn vị
    info1 = doc.add_paragraph()
    info1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_i1 = info1.add_run('Trung tâm Công nghệ thông tin')
    run_i1.font.size = Pt(14)
    run_i1.font.name = 'Times New Roman'

    info2 = doc.add_paragraph()
    info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_i2 = info2.add_run('Bộ Khoa học và Công nghệ')
    run_i2.bold = True
    run_i2.font.size = Pt(14)
    run_i2.font.name = 'Times New Roman'

    doc.add_paragraph()

    url_para = doc.add_paragraph()
    url_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_url = url_para.add_run('Địa chỉ truy cập: https://hientrangcds.mst.gov.vn')
    run_url.font.size = Pt(13)
    run_url.font.name = 'Times New Roman'

    ver_para = doc.add_paragraph()
    ver_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ver = ver_para.add_run('Phiên bản 1.0 – Tháng 01/2026')
    run_ver.font.size = Pt(12)
    run_ver.italic = True
    run_ver.font.name = 'Times New Roman'

    doc.add_page_break()

    # ==================== MỤC LỤC ====================
    doc.add_heading('Mục lục', 1)

    toc_items = [
        ('1.', 'Giới thiệu chung'),
        ('2.', 'Đăng nhập hệ thống'),
        ('3.', 'Hướng dẫn nhập liệu chi tiết'),
        ('   3.1.', 'Tab Thông tin cơ bản'),
        ('   3.2.', 'Tab Bối cảnh nghiệp vụ'),
        ('   3.3.', 'Tab Kiến trúc công nghệ'),
        ('   3.4.', 'Tab Kiến trúc dữ liệu'),
        ('   3.5.', 'Tab Tích hợp hệ thống'),
        ('   3.6.', 'Tab An toàn thông tin'),
        ('   3.7.', 'Tab Hạ tầng'),
        ('   3.8.', 'Tab Vận hành'),
        ('   3.9.', 'Tab Đánh giá'),
        ('4.', 'Những điểm cần lưu ý'),
        ('5.', 'Xử lý sự cố thường gặp'),
        ('6.', 'Thông tin liên hệ hỗ trợ'),
    ]

    for num, text in toc_items:
        p = doc.add_paragraph()
        run_num = p.add_run(num + ' ')
        run_num.font.name = 'Times New Roman'
        run_text = p.add_run(text)
        run_text.font.name = 'Times New Roman'

    doc.add_page_break()

    # ==================== 1. GIỚI THIỆU ====================
    doc.add_heading('1. Giới thiệu chung', 1)

    p1 = doc.add_paragraph(
        'Hệ thống Khảo sát Chuyển đổi số (CĐS) trực tuyến là công cụ được xây dựng '
        'nhằm hỗ trợ các đơn vị thuộc Bộ Khoa học và Công nghệ khai báo, theo dõi '
        'và quản lý thông tin về các hệ thống công nghệ thông tin đang vận hành.'
    )

    p2 = doc.add_paragraph(
        'Thông qua việc thu thập dữ liệu một cách có hệ thống, các đơn vị quản lý '
        'có thể đánh giá mức độ chuyển đổi số, phát hiện những điểm cần cải thiện, '
        'đồng thời lập kế hoạch phát triển công nghệ thông tin phù hợp với định hướng '
        'chung của Bộ.'
    )

    doc.add_heading('Mục đích sử dụng', 2)

    purposes = [
        'Thu thập thông tin đầy đủ về các hệ thống công nghệ thông tin đang hoạt động tại các đơn vị trực thuộc.',
        'Đánh giá mức độ chuyển đổi số của từng đơn vị cũng như toàn Bộ.',
        'Hỗ trợ việc lập kế hoạch nâng cấp, tích hợp và hiện đại hóa hệ thống.',
        'Cung cấp dữ liệu phục vụ báo cáo tổng hợp cho lãnh đạo.',
    ]

    for purpose in purposes:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(purpose)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    doc.add_heading('Đối tượng sử dụng', 2)

    p3 = doc.add_paragraph(
        'Tài liệu này dành cho cán bộ được phân công nhập liệu thông tin hệ thống '
        'tại các đơn vị. Người dùng không cần có kiến thức chuyên sâu về công nghệ '
        'thông tin, tuy nhiên cần nắm được thông tin cơ bản về các hệ thống đang '
        'vận hành tại đơn vị mình.'
    )

    doc.add_page_break()

    # ==================== 2. ĐĂNG NHẬP ====================
    doc.add_heading('2. Đăng nhập hệ thống', 1)

    p = doc.add_paragraph(
        'Để bắt đầu sử dụng hệ thống, anh/chị cần đăng nhập bằng tài khoản '
        'đã được quản trị viên cấp. Dưới đây là các bước thực hiện:'
    )

    doc.add_heading('Các bước đăng nhập', 2)

    login_steps = [
        ('Bước 1:', 'Mở trình duyệt web (Chrome, Firefox hoặc Edge) và truy cập địa chỉ:\nhttps://hientrangcds.mst.gov.vn'),
        ('Bước 2:', 'Tại màn hình đăng nhập, nhập Tên đăng nhập do quản trị viên cung cấp.'),
        ('Bước 3:', 'Nhập Mật khẩu tương ứng.'),
        ('Bước 4:', 'Nếu muốn hệ thống ghi nhớ phiên làm việc, tích vào ô "Ghi nhớ đăng nhập". Điều này giúp anh/chị không phải đăng nhập lại mỗi lần truy cập.'),
        ('Bước 5:', 'Nhấn nút "Đăng nhập" để vào hệ thống.'),
    ]

    for step_title, step_desc in login_steps:
        p = doc.add_paragraph()
        run_title = p.add_run(step_title + ' ')
        run_title.bold = True
        run_title.font.name = 'Times New Roman'
        run_desc = p.add_run(step_desc)
        run_desc.font.name = 'Times New Roman'

    doc.add_paragraph()
    add_note_box(doc,
        'Trường hợp quên mật khẩu, vui lòng liên hệ quản trị viên hệ thống để được '
        'cấp lại. Không nên tự thử nhiều lần vì tài khoản có thể bị khóa tạm thời.',
        "tip")

    doc.add_page_break()

    # ==================== 3. HƯỚNG DẪN NHẬP LIỆU ====================
    doc.add_heading('3. Hướng dẫn nhập liệu chi tiết', 1)

    # Khung cảnh báo quan trọng
    warning_para = doc.add_paragraph()
    run_warn = warning_para.add_run('⚠ NGUYÊN TẮC QUAN TRỌNG CẦN NHỚ')
    run_warn.bold = True
    run_warn.font.size = Pt(14)
    run_warn.font.name = 'Times New Roman'
    run_warn.font.color.rgb = RGBColor(192, 0, 0)  # Đỏ

    rule1 = doc.add_paragraph(
        'Hệ thống chia thông tin thành 9 tab (thẻ) riêng biệt. Mỗi tab chứa một nhóm '
        'thông tin liên quan. Anh/chị cần tuân thủ quy tắc sau:'
    )

    rules = [
        'Phải điền đầy đủ tất cả các trường có dấu sao (*) trong một tab trước khi chuyển sang tab tiếp theo.',
        'Nếu còn trường bắt buộc chưa điền, hệ thống sẽ hiển thị thông báo lỗi màu đỏ và không cho phép chuyển tab.',
        'Sau khi điền xong mỗi tab, nhấn nút "Lưu & Tiếp tục" để lưu dữ liệu và chuyển sang tab kế tiếp.',
        'Không thể bỏ qua bất kỳ tab nào – phải hoàn thành theo đúng thứ tự từ Tab 1 đến Tab 9.',
    ]

    for rule in rules:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(rule)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    doc.add_paragraph()

    # ==================== TAB 1 ====================
    doc.add_heading('3.1. Tab Thông tin cơ bản', 2)

    p = doc.add_paragraph(
        'Đây là tab đầu tiên, thu thập các thông tin nhận dạng cơ bản của hệ thống. '
        'Anh/chị cần hoàn thành tab này trước khi có thể điền các tab còn lại.'
    )

    add_image_with_caption(doc, f"{SCREENSHOTS_DIR}/tab1-co-ban-filled.png",
                          "Hình 1: Giao diện nhập liệu Tab Thông tin cơ bản")

    doc.add_heading('Giải thích từng trường:', 3)

    tab1_fields = [
        ('Tổ chức (*)',
         'Chọn đơn vị đang sở hữu hoặc quản lý hệ thống từ danh sách thả xuống. '
         'Danh sách này bao gồm tất cả các đơn vị trực thuộc Bộ.'),
        ('Mã hệ thống',
         'Mã định danh do hệ thống tự động tạo. Anh/chị không cần nhập trường này.'),
        ('Tên hệ thống (*)',
         'Nhập tên đầy đủ, chính thức của hệ thống bằng tiếng Việt. '
         'Ví dụ: "Hệ thống quản lý văn bản và điều hành".'),
        ('Tên tiếng Anh (*)',
         'Nhập tên hệ thống bằng tiếng Anh. Nếu hệ thống không có tên tiếng Anh, '
         'có thể nhập lại tên tiếng Việt hoặc phiên âm.'),
        ('Mô tả (*)',
         'Viết mô tả ngắn gọn về chức năng chính và mục đích của hệ thống. '
         'Nên viết từ 2-3 câu để người đọc hiểu được hệ thống dùng để làm gì.'),
        ('Trạng thái (*)',
         'Chọn trạng thái hiện tại của hệ thống: Đang vận hành (hoạt động bình thường), '
         'Thí điểm (đang trong giai đoạn thử nghiệm), Bảo trì (tạm ngừng để nâng cấp), v.v.'),
        ('Mức độ quan trọng (*)',
         'Đánh giá tầm quan trọng của hệ thống đối với hoạt động của đơn vị: '
         'Cực kỳ quan trọng, Quan trọng, Trung bình hoặc Thấp.'),
        ('Phạm vi sử dụng (*)',
         'Chọn "Nội bộ" nếu chỉ cán bộ trong đơn vị sử dụng. '
         'Chọn "Bên ngoài" nếu có người dân hoặc doanh nghiệp truy cập sử dụng.'),
        ('Nhu cầu (*)',
         'Xác định nhu cầu phát triển: Cần tích hợp liên thông với hệ thống khác, '
         'Cần nâng cấp, Cần thay thế bằng hệ thống mới, hoặc Duy trì nguyên trạng.'),
        ('Thời gian mong muốn hoàn thành (*)',
         'Chọn tháng và năm dự kiến hoàn thành nếu có kế hoạch nâng cấp/thay thế.'),
        ('Nhóm hệ thống (*)',
         'Phân loại hệ thống: Cơ sở dữ liệu chuyên ngành, Ứng dụng nội bộ, '
         'Cổng thông tin điện tử, Dịch vụ công trực tuyến, v.v.'),
        ('Ghi chú bổ sung (*)',
         'Thông tin thêm cho tab này. Nếu không có gì bổ sung, nhập "Không có".'),
    ]

    for field, desc in tab1_fields:
        p = doc.add_paragraph()
        run_field = p.add_run(f'• {field}: ')
        run_field.bold = True
        run_field.font.name = 'Times New Roman'
        run_desc = p.add_run(desc)
        run_desc.font.name = 'Times New Roman'

    doc.add_page_break()

    # ==================== TAB 2 ====================
    doc.add_heading('3.2. Tab Bối cảnh nghiệp vụ', 2)

    p = doc.add_paragraph(
        'Tab này ghi nhận thông tin về mục tiêu nghiệp vụ mà hệ thống phục vụ, '
        'đối tượng người dùng và các chỉ số về quy mô sử dụng.'
    )

    add_image_with_caption(doc, f"{SCREENSHOTS_DIR}/tab2-nghiep-vu-filled.png",
                          "Hình 2: Giao diện nhập liệu Tab Bối cảnh nghiệp vụ")

    doc.add_heading('Giải thích từng trường:', 3)

    tab2_fields = [
        ('Mục tiêu nghiệp vụ (*)',
         'Chọn một hoặc nhiều mục tiêu mà hệ thống hướng tới: Số hóa quy trình công việc, '
         'Cải thiện chất lượng dịch vụ công, Tăng cường minh bạch, Hỗ trợ ra quyết định, v.v.'),
        ('Quy trình nghiệp vụ chính (*)',
         'Chọn các quy trình mà hệ thống hỗ trợ thực hiện: Quản lý hồ sơ, Phê duyệt, '
         'Tra cứu thông tin, Lập báo cáo, Thống kê số liệu, v.v.'),
        ('Có đủ hồ sơ phân tích thiết kế? (*)',
         'Bật công tắc nếu hệ thống có đầy đủ tài liệu phân tích, thiết kế kỹ thuật.'),
        ('Đối tượng sử dụng (*)',
         'Chọn các nhóm người dùng: Cán bộ nội bộ, Doanh nghiệp, Người dân, '
         'Cơ quan địa phương, Đối tác, v.v.'),
        ('Số lượng người dùng hàng năm (*)',
         'Ước tính số người sử dụng hệ thống trong một năm.'),
        ('Tổng số tài khoản (*)',
         'Tổng số tài khoản đã đăng ký trên hệ thống, bao gồm cả tài khoản đang hoạt động '
         'và không còn hoạt động.'),
        ('Số người dùng hoạt động hàng tháng (*)',
         'Viết tắt MAU - Monthly Active Users. Số người dùng truy cập ít nhất một lần trong tháng.'),
        ('Số người dùng hoạt động hàng ngày (*)',
         'Viết tắt DAU - Daily Active Users. Số người dùng truy cập trung bình mỗi ngày.'),
        ('Số đơn vị/địa phương sử dụng (*)',
         'Số lượng cơ quan, đơn vị hoặc địa phương đang sử dụng hệ thống.'),
        ('Ghi chú bổ sung (*)',
         'Thông tin thêm về nghiệp vụ. Nếu không có, nhập "Không có".'),
    ]

    for field, desc in tab2_fields:
        p = doc.add_paragraph()
        run_field = p.add_run(f'• {field}: ')
        run_field.bold = True
        run_field.font.name = 'Times New Roman'
        run_desc = p.add_run(desc)
        run_desc.font.name = 'Times New Roman'

    doc.add_page_break()

    # ==================== TAB 3 ====================
    doc.add_heading('3.3. Tab Kiến trúc công nghệ', 2)

    p = doc.add_paragraph(
        'Tab này yêu cầu thông tin chi tiết về các công nghệ được sử dụng trong hệ thống. '
        'Nếu anh/chị không nắm rõ các thông tin kỹ thuật, hãy liên hệ bộ phận phát triển '
        'hoặc vận hành để được hỗ trợ.'
    )

    add_image_with_caption(doc, f"{SCREENSHOTS_DIR}/tab3-cong-nghe-filled.png",
                          "Hình 3: Giao diện nhập liệu Tab Kiến trúc công nghệ")

    doc.add_heading('Giải thích từng trường:', 3)

    tab3_fields = [
        ('Ngôn ngữ lập trình (*)',
         'Chọn ngôn ngữ được sử dụng để phát triển hệ thống: Python, Java, JavaScript, C#, PHP, v.v. '
         'Có thể chọn nhiều ngôn ngữ.'),
        ('Framework/Thư viện (*)',
         'Chọn các framework (khung phát triển) đang dùng: Django, Spring Boot, React, Angular, Vue.js, Laravel, v.v.'),
        ('Cơ sở dữ liệu (*)',
         'Chọn hệ quản trị cơ sở dữ liệu: SQL Server, MySQL, PostgreSQL, Oracle, MongoDB, v.v.'),
        ('Nền tảng triển khai (*)',
         'Chọn nơi hệ thống được triển khai: Máy chủ nội bộ (On-premise), AWS, Azure, Google Cloud, v.v.'),
        ('Công nghệ phía máy chủ (*)',
         'Công nghệ xử lý phía server: Node.js, Python, Java, C#/.NET, PHP, v.v.'),
        ('Công nghệ giao diện (*)',
         'Công nghệ phát triển giao diện người dùng: React, Vue.js, Angular, HTML/CSS thuần, v.v.'),
        ('Loại kiến trúc (*)',
         'Chọn kiến trúc hệ thống: Monolithic (một khối), Microservices (vi dịch vụ), Serverless, SaaS, v.v.'),
        ('Container hóa (*)',
         'Công nghệ container: Docker, Kubernetes, OpenShift. Chọn "Không sử dụng" nếu không áp dụng.'),
        ('Hỗ trợ nhiều đơn vị thuê chung (*)',
         'Hệ thống có cho phép nhiều đơn vị dùng chung một bản cài đặt không (Multi-tenant)?'),
        ('Kiến trúc phân lớp (*)',
         'Hệ thống có chia thành các lớp riêng biệt (Presentation, Business Logic, Data Access) không?'),
        ('Chi tiết phân lớp (*)',
         'Mô tả cụ thể các lớp nếu có kiến trúc phân lớp.'),
        ('Kiểu giao tiếp API (*)',
         'Chọn: REST API, GraphQL, gRPC, SOAP hoặc ghi rõ nếu dùng loại khác.'),
        ('Hệ thống hàng đợi (*)',
         'Chọn: Apache Kafka, RabbitMQ, ActiveMQ. Chọn "Không sử dụng" nếu không áp dụng.'),
        ('Hệ thống lưu đệm (*)',
         'Chọn: Redis, Memcached. Chọn "Không sử dụng" nếu không áp dụng.'),
        ('Công cụ tìm kiếm (*)',
         'Chọn: Elasticsearch, Solr. Chọn "Không sử dụng" nếu không áp dụng.'),
        ('Công cụ báo cáo (*)',
         'Chọn: Power BI, Tableau, Tự phát triển, Không có.'),
        ('Kho mã nguồn (*)',
         'Chọn: GitHub, GitLab, Bitbucket. Chọn "Không quản lý" nếu mã nguồn không được lưu trữ có hệ thống.'),
        ('Có quy trình tích hợp liên tục? (*)',
         'Hệ thống có thiết lập CI/CD (tự động build, test, deploy) không?'),
        ('Công cụ CI/CD (*)',
         'Chọn: Jenkins, GitHub Actions, GitLab CI, Azure DevOps, v.v.'),
        ('Có kiểm thử tự động? (*)',
         'Hệ thống có viết và chạy test tự động không?'),
        ('Công cụ kiểm thử (*)',
         'Liệt kê công cụ: Jest, Pytest, Selenium, JUnit, v.v.'),
        ('Ghi chú bổ sung (*)',
         'Thông tin bổ sung về công nghệ. Nếu không có, nhập "Không có".'),
    ]

    for field, desc in tab3_fields:
        p = doc.add_paragraph()
        run_field = p.add_run(f'• {field}: ')
        run_field.bold = True
        run_field.font.name = 'Times New Roman'
        run_desc = p.add_run(desc)
        run_desc.font.name = 'Times New Roman'

    doc.add_page_break()

    # ==================== TAB 4 ====================
    doc.add_heading('3.4. Tab Kiến trúc dữ liệu', 2)

    p = doc.add_paragraph(
        'Tab này thu thập thông tin về cách dữ liệu được quản lý, lưu trữ và phân loại '
        'trong hệ thống.'
    )

    add_image_with_caption(doc, f"{SCREENSHOTS_DIR}/tab4-du-lieu-filled.png",
                          "Hình 4: Giao diện nhập liệu Tab Kiến trúc dữ liệu")

    doc.add_heading('Giải thích từng trường:', 3)

    tab4_fields = [
        ('Nguồn dữ liệu (*)',
         'Chọn nguồn dữ liệu đầu vào: Người dùng nhập liệu, Kết nối API bên ngoài, '
         'Đồng bộ từ hệ thống khác, Nhập từ file, v.v.'),
        ('Loại dữ liệu (*)',
         'Chọn loại dữ liệu chính: Dữ liệu nghiệp vụ, Văn bản tài liệu, Số liệu thống kê, '
         'Dữ liệu danh mục dùng chung, v.v.'),
        ('Phân loại mức độ bảo mật (*)',
         'Chọn theo quy định: Công khai, Nội bộ, Hạn chế, Bí mật hoặc Tối mật.'),
        ('Khối lượng dữ liệu (*)',
         'Ước tính tổng dung lượng: Dưới 1GB, 1-10GB, 10-100GB, 100GB-1TB, hoặc Trên 1TB.'),
        ('Dung lượng cơ sở dữ liệu hiện tại (*)',
         'Nhập số GB dung lượng thực tế của database.'),
        ('Dung lượng file đính kèm (*)',
         'Nhập số GB dung lượng các file văn bản, hình ảnh đính kèm.'),
        ('Tốc độ tăng trưởng dữ liệu (*)',
         'Ước tính phần trăm tăng trưởng dữ liệu hàng năm.'),
        ('Loại lưu trữ file (*)',
         'Chọn: File Server, Object Storage (S3), NAS, Lưu trong database, v.v.'),
        ('Số bản ghi (*)',
         'Ước tính tổng số bản ghi (record) trong cơ sở dữ liệu.'),
        ('Cơ sở dữ liệu phụ (*)',
         'Liệt kê các CSDL phụ hoặc cache nếu có: MySQL phụ, Redis cache, MongoDB, v.v.'),
        ('Chính sách lưu trữ',
         'Mô tả quy định về thời gian lưu trữ, sao lưu, và xóa dữ liệu.'),
        ('Có danh mục dữ liệu? (*)',
         'Hệ thống có duy trì danh mục (Data Catalog) mô tả các trường dữ liệu không?'),
        ('Có quản lý dữ liệu chủ? (*)',
         'Hệ thống có sử dụng Master Data Management để đảm bảo tính nhất quán dữ liệu không?'),
    ]

    for field, desc in tab4_fields:
        p = doc.add_paragraph()
        run_field = p.add_run(f'• {field}: ')
        run_field.bold = True
        run_field.font.name = 'Times New Roman'
        run_desc = p.add_run(desc)
        run_desc.font.name = 'Times New Roman'

    doc.add_page_break()

    # ==================== TAB 5 ====================
    doc.add_heading('3.5. Tab Tích hợp hệ thống', 2)

    p = doc.add_paragraph(
        'Tab này ghi nhận thông tin về cách hệ thống kết nối và trao đổi dữ liệu '
        'với các hệ thống khác trong và ngoài đơn vị.'
    )

    add_image_with_caption(doc, f"{SCREENSHOTS_DIR}/tab5-tich-hop-filled.png",
                          "Hình 5: Giao diện nhập liệu Tab Tích hợp hệ thống")

    doc.add_heading('Giải thích từng trường:', 3)

    tab5_fields = [
        ('Số API cung cấp (*)',
         'Số lượng API (giao diện lập trình) mà hệ thống này cung cấp cho các hệ thống khác sử dụng.'),
        ('Số API tiêu thụ (*)',
         'Số lượng API từ các hệ thống khác mà hệ thống này đang gọi đến.'),
        ('Chuẩn API (*)',
         'Chọn chuẩn mô tả API: OpenAPI 3.0, OpenAPI 2.0 (Swagger), SOAP WSDL, GraphQL, gRPC, v.v.'),
        ('Có API Gateway? (*)',
         'Hệ thống có sử dụng cổng API tập trung để quản lý các API không?'),
        ('Tên API Gateway (*)',
         'Chọn sản phẩm: Kong, AWS API Gateway, Azure API Management, Apigee. '
         'Chọn "Không có" nếu không sử dụng.'),
        ('Có quản lý phiên bản API? (*)',
         'API có được quản lý theo phiên bản (v1, v2, v3...) không?'),
        ('Có giới hạn tần suất gọi? (*)',
         'Có thiết lập Rate Limiting để hạn chế số lượng request trong khoảng thời gian không?'),
        ('Tài liệu API (*)',
         'Nhập đường link hoặc mô tả về tài liệu hướng dẫn sử dụng API (Swagger/OpenAPI docs).'),
        ('Chuẩn đánh số phiên bản (*)',
         'Chọn cách đánh phiên bản: Trong đường dẫn URL, Trong Header, Trong Query Parameter, v.v.'),
        ('Có giám sát tích hợp? (*)',
         'Có hệ thống theo dõi trạng thái các kết nối tích hợp không?'),
        ('Hệ thống nội bộ tích hợp (*)',
         'Chọn các hệ thống nội bộ đang kết nối: Quản lý văn bản, Nhân sự, Tài chính kế toán, '
         'Cổng thông tin nội bộ, v.v.'),
        ('Hệ thống bên ngoài tích hợp (*)',
         'Chọn các hệ thống bên ngoài: VNeID, LGSP, Cổng Dịch vụ công quốc gia, '
         'Hệ thống Thuế, Hải quan, v.v.'),
        ('API/Webservices (*)',
         'Liệt kê các endpoint API hoặc webservice chính.'),
        ('Phương thức trao đổi dữ liệu (*)',
         'Chọn: Gọi API REST, Gọi API SOAP, Truyền file, Kết nối trực tiếp Database, v.v.'),
    ]

    for field, desc in tab5_fields:
        p = doc.add_paragraph()
        run_field = p.add_run(f'• {field}: ')
        run_field.bold = True
        run_field.font.name = 'Times New Roman'
        run_desc = p.add_run(desc)
        run_desc.font.name = 'Times New Roman'

    doc.add_page_break()

    # ==================== TAB 6 ====================
    doc.add_heading('3.6. Tab An toàn thông tin', 2)

    p = doc.add_paragraph(
        'Tab này thu thập thông tin về các biện pháp bảo mật và đảm bảo an toàn thông tin '
        'đang được áp dụng cho hệ thống.'
    )

    doc.add_heading('Giải thích từng trường:', 3)

    tab6_fields = [
        ('Phương thức xác thực (*)',
         'Chọn cách người dùng đăng nhập: Tên đăng nhập/Mật khẩu, Đăng nhập một lần (SSO), '
         'Xác thực hai bước (2FA), OAuth, Kết nối LDAP, v.v.'),
        ('Có mã hóa dữ liệu? (*)',
         'Hệ thống có mã hóa các dữ liệu nhạy cảm như mật khẩu, thông tin cá nhân không?'),
        ('Có ghi nhật ký hoạt động? (*)',
         'Hệ thống có lưu lại log các thao tác của người dùng để kiểm tra sau này không?'),
        ('Mức độ an toàn thông tin (*)',
         'Chọn cấp độ theo quy định: Cấp 1 (thấp nhất) đến Cấp 5 (cao nhất).'),
        ('Có tài liệu an toàn thông tin? (*)',
         'Hệ thống có đầy đủ hồ sơ, tài liệu về an toàn thông tin theo quy định không?'),
        ('Ghi chú bổ sung (*)',
         'Thông tin thêm về bảo mật. Nếu không có, nhập "Không có".'),
    ]

    for field, desc in tab6_fields:
        p = doc.add_paragraph()
        run_field = p.add_run(f'• {field}: ')
        run_field.bold = True
        run_field.font.name = 'Times New Roman'
        run_desc = p.add_run(desc)
        run_desc.font.name = 'Times New Roman'

    doc.add_page_break()

    # ==================== TAB 7 ====================
    doc.add_heading('3.7. Tab Hạ tầng', 2)

    p = doc.add_paragraph(
        'Tab này ghi nhận thông tin về hạ tầng phần cứng, máy chủ và kế hoạch '
        'dự phòng phục vụ vận hành hệ thống.'
    )

    doc.add_heading('Giải thích từng trường:', 3)

    tab7_fields = [
        ('Cấu hình máy chủ (*)',
         'Mô tả cấu hình: Số lượng CPU, dung lượng RAM, dung lượng ổ cứng (Storage).'),
        ('Phương án sao lưu (*)',
         'Chọn: Sao lưu hàng ngày, Sao lưu hàng tuần, Sao lưu theo thời gian thực, v.v.'),
        ('Dung lượng lưu trữ (*)',
         'Tổng dung lượng storage được cấp phát cho hệ thống.'),
        ('Kế hoạch phục hồi thảm họa (*)',
         'Mô tả kế hoạch khôi phục hệ thống khi có sự cố nghiêm trọng (DR Plan).'),
        ('Vị trí triển khai (*)',
         'Chọn: Trung tâm dữ liệu của Bộ, Thuê chỗ đặt máy chủ, Dịch vụ Cloud, v.v.'),
        ('Cấu hình tính toán (*)',
         'Chi tiết về máy chủ ảo (VM), Container hoặc máy chủ vật lý đang sử dụng.'),
        ('Loại hạ tầng tính toán (*)',
         'Chọn: Máy ảo (VM), Container, Máy chủ vật lý, Serverless.'),
        ('Tần suất triển khai (*)',
         'Chọn tần suất cập nhật phiên bản: Hàng ngày, Hàng tuần, Hàng tháng, hoặc Khi cần thiết.'),
        ('Ghi chú bổ sung (*)',
         'Thông tin thêm về hạ tầng. Nếu không có, nhập "Không có".'),
    ]

    for field, desc in tab7_fields:
        p = doc.add_paragraph()
        run_field = p.add_run(f'• {field}: ')
        run_field.bold = True
        run_field.font.name = 'Times New Roman'
        run_desc = p.add_run(desc)
        run_desc.font.name = 'Times New Roman'

    doc.add_page_break()

    # ==================== TAB 8 ====================
    doc.add_heading('3.8. Tab Vận hành', 2)

    p = doc.add_paragraph(
        'Tab này ghi nhận thông tin về đội ngũ phụ trách và cách thức hỗ trợ '
        'vận hành hệ thống.'
    )

    doc.add_heading('Giải thích từng trường:', 3)

    tab8_fields = [
        ('Người phụ trách nghiệp vụ (*)',
         'Họ tên cán bộ chịu trách nhiệm về mặt nghiệp vụ, đảm bảo hệ thống đáp ứng yêu cầu công việc.'),
        ('Người phụ trách kỹ thuật (*)',
         'Họ tên cán bộ chịu trách nhiệm về mặt kỹ thuật, xử lý sự cố và bảo trì hệ thống.'),
        ('Người chịu trách nhiệm chính (*)',
         'Họ tên người liên hệ chính khi cần hỗ trợ hoặc có vấn đề về hệ thống.'),
        ('Số điện thoại liên hệ (*)',
         'Số điện thoại để liên hệ khi cần hỗ trợ khẩn cấp.'),
        ('Email liên hệ (*)',
         'Địa chỉ email để gửi yêu cầu hỗ trợ.'),
        ('Mức độ hỗ trợ (*)',
         'Chọn: Hỗ trợ 24/7, Trong giờ hành chính, Theo yêu cầu, v.v.'),
        ('Ghi chú bổ sung (*)',
         'Thông tin thêm về vận hành. Nếu không có, nhập "Không có".'),
    ]

    for field, desc in tab8_fields:
        p = doc.add_paragraph()
        run_field = p.add_run(f'• {field}: ')
        run_field.bold = True
        run_field.font.name = 'Times New Roman'
        run_desc = p.add_run(desc)
        run_desc.font.name = 'Times New Roman'

    doc.add_page_break()

    # ==================== TAB 9 ====================
    doc.add_heading('3.9. Tab Đánh giá', 2)

    p = doc.add_paragraph(
        'Tab cuối cùng này thu thập các đánh giá về hiệu năng, chất lượng hiện tại '
        'và kế hoạch phát triển trong tương lai của hệ thống.'
    )

    doc.add_heading('Giải thích từng trường:', 3)

    tab9_fields = [
        ('Đánh giá hiệu năng (*)',
         'Chọn mức đánh giá: Rất tốt, Tốt, Trung bình, Kém hoặc Rất kém.'),
        ('Đánh giá mức độ hài lòng (*)',
         'Đánh giá sự hài lòng của người dùng: Rất hài lòng, Hài lòng, Bình thường, Không hài lòng.'),
        ('Mức độ nợ kỹ thuật (*)',
         'Đánh giá mức độ code cũ, cần refactor: Thấp, Trung bình, Cao hoặc Rất cao.'),
        ('Đề xuất hành động (*)',
         'Chọn hướng phát triển: Duy trì nguyên trạng, Cần nâng cấp, Cần thay thế, hoặc Nên loại bỏ.'),
        ('Điểm mạnh về tích hợp (*)',
         'Chọn các điểm mạnh của hệ thống trong việc tích hợp với hệ thống khác.'),
        ('Các vướng mắc (*)',
         'Chọn các khó khăn, vấn đề đang gặp phải.'),
        ('Tỷ lệ thời gian hoạt động (*)',
         'Nhập phần trăm uptime, ví dụ: 99.9 nghĩa là hệ thống hoạt động 99.9% thời gian.'),
        ('Thời gian phản hồi trung bình (*)',
         'Nhập thời gian response trung bình tính bằng mili-giây (ms).'),
        ('Kế hoạch thay thế (*)',
         'Mô tả kế hoạch thay thế hệ thống nếu có. Nếu không, nhập "Không có".'),
        ('Các vấn đề chính (*)',
         'Liệt kê các vấn đề nghiêm trọng đang gặp. Nếu không có, nhập "Không có".'),
        ('Đề xuất cải tiến (*)',
         'Ghi các đề xuất để nâng cao chất lượng, hiệu năng của hệ thống.'),
        ('Kế hoạch tương lai (*)',
         'Mô tả định hướng phát triển trong thời gian tới.'),
        ('Mức độ ưu tiên hiện đại hóa (*)',
         'Đánh giá mức ưu tiên để nâng cấp, hiện đại hóa: Cao, Trung bình hoặc Thấp.'),
    ]

    for field, desc in tab9_fields:
        p = doc.add_paragraph()
        run_field = p.add_run(f'• {field}: ')
        run_field.bold = True
        run_field.font.name = 'Times New Roman'
        run_desc = p.add_run(desc)
        run_desc.font.name = 'Times New Roman'

    doc.add_page_break()

    # ==================== 4. LƯU Ý QUAN TRỌNG ====================
    doc.add_heading('4. Những điểm cần lưu ý', 1)

    p = doc.add_paragraph(
        'Dưới đây là một số lưu ý quan trọng giúp anh/chị hoàn thành việc nhập liệu '
        'một cách thuận lợi và chính xác:'
    )

    important_notes = [
        ('Điền đầy đủ các trường bắt buộc',
         'Mỗi tab đều có các trường đánh dấu sao (*). Hệ thống sẽ không cho phép '
         'chuyển sang tab tiếp theo hoặc lưu dữ liệu nếu còn bất kỳ trường bắt buộc nào chưa được điền.'),

        ('Tuân thủ thứ tự các tab',
         'Anh/chị cần hoàn thành từng tab theo đúng thứ tự từ 1 đến 9. '
         'Việc này đảm bảo dữ liệu được ghi nhận đầy đủ và có hệ thống.'),

        ('Lưu dữ liệu thường xuyên',
         'Sau mỗi tab, hãy nhấn nút "Lưu & Tiếp tục". Điều này giúp tránh mất dữ liệu '
         'trong trường hợp mất kết nối hoặc hết phiên làm việc.'),

        ('Khi không biết thông tin chính xác',
         'Nếu chưa nắm rõ thông tin, anh/chị có thể nhập giá trị ước tính, '
         'hoặc ghi "Không có" / "Chưa xác định" tùy loại trường. '
         'Quan trọng là không để trống các trường bắt buộc.'),

        ('Liên hệ bộ phận kỹ thuật khi cần',
         'Với các tab về công nghệ (Tab 3, 4, 5), nếu không chắc chắn, '
         'hãy hỏi đội ngũ phát triển hoặc vận hành hệ thống để có thông tin chính xác.'),

        ('Kiểm tra lại trước khi hoàn thành',
         'Sau khi điền xong tất cả 9 tab, nên dành vài phút rà soát lại thông tin '
         'để đảm bảo không có sai sót.'),
    ]

    for title, desc in important_notes:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run_title = p.add_run(f'✓ {title}')
        run_title.bold = True
        run_title.font.name = 'Times New Roman'

        p2 = doc.add_paragraph(desc)
        p2.paragraph_format.left_indent = Cm(0.5)

    doc.add_page_break()

    # ==================== 5. XỬ LÝ SỰ CỐ ====================
    doc.add_heading('5. Xử lý sự cố thường gặp', 1)

    p = doc.add_paragraph(
        'Trong quá trình sử dụng, anh/chị có thể gặp một số vấn đề sau. '
        'Dưới đây là cách xử lý:'
    )

    errors = [
        ('Thông báo "Vui lòng điền đủ X trường bắt buộc..."',
         'Nguyên nhân: Còn trường bắt buộc chưa được điền trong tab hiện tại.',
         'Cách xử lý: Cuộn lại từ đầu tab, tìm các trường có viền đỏ và thông báo lỗi phía dưới, '
         'sau đó điền đầy đủ thông tin vào các trường đó.'),

        ('Không thể chuyển sang tab tiếp theo',
         'Nguyên nhân: Tab hiện tại chưa hoàn thành.',
         'Cách xử lý: Kiểm tra và hoàn thành tất cả các trường bắt buộc trong tab đang mở, '
         'sau đó nhấn "Lưu & Tiếp tục".'),

        ('Nút "Lưu" bị mờ, không nhấn được',
         'Nguyên nhân: Dữ liệu chưa có thay đổi so với lần lưu trước hoặc form chưa hợp lệ.',
         'Cách xử lý: Kiểm tra xem đã điền đủ các trường bắt buộc chưa. '
         'Nếu đã đủ mà nút vẫn mờ, thử làm mới trang (nhấn F5) rồi đăng nhập lại.'),

        ('Phiên làm việc hết hạn',
         'Nguyên nhân: Không thao tác trên hệ thống trong thời gian dài.',
         'Cách xử lý: Đăng nhập lại bằng tài khoản của mình. '
         'Để tránh tình trạng này, có thể tích vào ô "Ghi nhớ đăng nhập" khi đăng nhập.'),

        ('Không lưu được dữ liệu',
         'Nguyên nhân: Có thể do mất kết nối mạng hoặc máy chủ đang bảo trì.',
         'Cách xử lý: Chờ vài phút rồi thử lại. Nếu vấn đề kéo dài, liên hệ quản trị viên.'),
    ]

    for error_title, cause, solution in errors:
        doc.add_heading(error_title, 2)

        p1 = doc.add_paragraph()
        run_c = p1.add_run('Nguyên nhân: ')
        run_c.bold = True
        run_c.font.name = 'Times New Roman'
        run_c_text = p1.add_run(cause.replace('Nguyên nhân: ', ''))
        run_c_text.font.name = 'Times New Roman'

        p2 = doc.add_paragraph()
        run_s = p2.add_run('Cách xử lý: ')
        run_s.bold = True
        run_s.font.name = 'Times New Roman'
        run_s_text = p2.add_run(solution.replace('Cách xử lý: ', ''))
        run_s_text.font.name = 'Times New Roman'

    doc.add_page_break()

    # ==================== 6. LIÊN HỆ HỖ TRỢ ====================
    doc.add_heading('6. Thông tin liên hệ hỗ trợ', 1)

    p = doc.add_paragraph(
        'Khi gặp khó khăn trong quá trình sử dụng hệ thống, anh/chị có thể liên hệ '
        'theo thông tin sau để được hỗ trợ:'
    )

    doc.add_paragraph()

    contact_info = [
        ('Đơn vị hỗ trợ:', 'Trung tâm Công nghệ thông tin - Bộ Khoa học và Công nghệ'),
        ('Địa chỉ:', '113 Trần Duy Hưng, Cầu Giấy, Hà Nội'),
        ('Email:', 'support@mst.gov.vn'),
        ('Điện thoại:', '024-xxx-xxxx'),
        ('Thời gian hỗ trợ:', 'Từ 8:00 đến 17:00, các ngày làm việc trong tuần'),
    ]

    for label, value in contact_info:
        p = doc.add_paragraph()
        run_label = p.add_run(label + ' ')
        run_label.bold = True
        run_label.font.name = 'Times New Roman'
        run_value = p.add_run(value)
        run_value.font.name = 'Times New Roman'

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    footer_line = doc.add_paragraph()
    footer_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = footer_line.add_run('─' * 30)
    run_line.font.name = 'Times New Roman'

    footer_note = doc.add_paragraph()
    footer_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fn = footer_note.add_run('Tài liệu được cập nhật lần cuối: Tháng 01/2026')
    run_fn.italic = True
    run_fn.font.size = Pt(12)
    run_fn.font.name = 'Times New Roman'

    version_note = doc.add_paragraph()
    version_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_vn = version_note.add_run('Phiên bản 1.0 – Humanized Edition')
    run_vn.italic = True
    run_vn.font.size = Pt(11)
    run_vn.font.name = 'Times New Roman'

    # Lưu file
    output_path = "Huong_Dan_Su_Dung_He_Thong_CDS_Humanized.docx"
    doc.save(output_path)
    print(f"✓ Tài liệu đã được tạo thành công: {output_path}")
    print(f"  - Font: Times New Roman 14pt")
    print(f"  - Định dạng: Theo chuẩn Nghị định 30/2020/NĐ-CP")
    print(f"  - Văn phong: Tự nhiên, thân thiện với người dùng")
    return output_path

if __name__ == "__main__":
    create_user_guide()
