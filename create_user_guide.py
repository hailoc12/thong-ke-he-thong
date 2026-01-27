#!/usr/bin/env python3
"""
Script to create User Guide DOCX for Hệ thống Khảo sát CĐS
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Path to screenshots
SCREENSHOTS_DIR = ".playwright-mcp/screenshots"

def add_heading_style(doc):
    """Add custom heading styles"""
    styles = doc.styles

    # Configure existing Heading 1
    heading1_style = styles['Heading 1']
    heading1_style.font.size = Pt(18)
    heading1_style.font.bold = True
    heading1_style.font.color.rgb = None  # Use theme color

    # Configure existing Heading 2
    heading2_style = styles['Heading 2']
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True

def add_image_with_caption(doc, image_path, caption, width=Inches(6)):
    """Add image with caption"""
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add caption
        caption_para = doc.add_paragraph(caption)
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_para.runs[0].italic = True
        caption_para.runs[0].font.size = Pt(10)
    else:
        doc.add_paragraph(f"[Hình ảnh: {caption}]")

def create_user_guide():
    """Create the user guide document"""
    doc = Document()
    add_heading_style(doc)

    # ==================== TITLE PAGE ====================
    title = doc.add_heading('HƯỚNG DẪN SỬ DỤNG', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('HỆ THỐNG KHẢO SÁT CHUYỂN ĐỔI SỐ TRỰC TUYẾN')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.bold = True

    doc.add_paragraph()

    info = doc.add_paragraph('Trung tâm Công nghệ thông tin - Bộ Khoa học và Công nghệ')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph('URL: https://hientrangcds.mst.gov.vn')
    doc.add_paragraph('Phiên bản: 1.0.0 - 2026')

    doc.add_page_break()

    # ==================== MỤC LỤC ====================
    doc.add_heading('MỤC LỤC', 1)

    toc_items = [
        '1. Giới thiệu chung',
        '2. Đăng nhập hệ thống',
        '3. Hướng dẫn nhập liệu',
        '   3.1. Tab 1: Thông tin cơ bản',
        '   3.2. Tab 2: Bối cảnh nghiệp vụ',
        '   3.3. Tab 3: Kiến trúc công nghệ',
        '   3.4. Tab 4: Kiến trúc dữ liệu',
        '   3.5. Tab 5: Tích hợp hệ thống',
        '   3.6. Tab 6: An toàn thông tin',
        '   3.7. Tab 7: Hạ tầng',
        '   3.8. Tab 8: Vận hành',
        '   3.9. Tab 9: Đánh giá',
        '4. Lưu ý quan trọng',
        '5. Xử lý lỗi thường gặp',
    ]

    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    # ==================== 1. GIỚI THIỆU ====================
    doc.add_heading('1. GIỚI THIỆU CHUNG', 1)

    doc.add_paragraph(
        'Hệ thống Khảo sát Chuyển đổi số (CĐS) trực tuyến là công cụ giúp các đơn vị '
        'thuộc Bộ Khoa học và Công nghệ khai báo, quản lý thông tin về các hệ thống CNTT '
        'đang vận hành. Dữ liệu thu thập được sử dụng để đánh giá mức độ chuyển đổi số '
        'và lập kế hoạch phát triển CNTT của Bộ.'
    )

    doc.add_heading('Mục đích của hệ thống', 2)
    purposes = [
        'Thu thập thông tin về các hệ thống CNTT đang vận hành tại các đơn vị',
        'Đánh giá mức độ chuyển đổi số của từng đơn vị và toàn Bộ',
        'Hỗ trợ lập kế hoạch nâng cấp, tích hợp, hiện đại hóa hệ thống',
        'Tạo báo cáo tổng hợp phục vụ công tác quản lý',
    ]
    for p in purposes:
        doc.add_paragraph(f'• {p}')

    doc.add_page_break()

    # ==================== 2. ĐĂNG NHẬP ====================
    doc.add_heading('2. ĐĂNG NHẬP HỆ THỐNG', 1)

    doc.add_paragraph('Để sử dụng hệ thống, người dùng cần đăng nhập với tài khoản được cấp.')

    doc.add_heading('Các bước đăng nhập:', 2)
    login_steps = [
        'Bước 1: Truy cập địa chỉ https://hientrangcds.mst.gov.vn',
        'Bước 2: Nhập Tên đăng nhập (do quản trị viên cung cấp)',
        'Bước 3: Nhập Mật khẩu',
        'Bước 4: (Tùy chọn) Tích vào ô "Ghi nhớ đăng nhập" nếu muốn duy trì phiên làm việc',
        'Bước 5: Nhấn nút "Đăng nhập"',
    ]
    for step in login_steps:
        doc.add_paragraph(step)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Lưu ý: ').bold = True
    p.add_run('Nếu quên mật khẩu, vui lòng liên hệ quản trị viên hệ thống để được hỗ trợ.')

    doc.add_page_break()

    # ==================== 3. HƯỚNG DẪN NHẬP LIỆU ====================
    doc.add_heading('3. HƯỚNG DẪN NHẬP LIỆU', 1)

    # Important note about required fields
    important_box = doc.add_paragraph()
    important_box.add_run('⚠️ LƯU Ý QUAN TRỌNG: ').bold = True

    doc.add_paragraph(
        'Hệ thống yêu cầu PHẢI NHẬP ĐẦY ĐỦ TẤT CẢ các trường có dấu (*) trong MỘT TAB '
        'trước khi được phép chuyển sang tab tiếp theo hoặc lưu dữ liệu.'
    )

    doc.add_paragraph(
        'Nếu còn trường bắt buộc chưa điền, hệ thống sẽ hiển thị thông báo lỗi màu đỏ '
        'và không cho phép chuyển tab.'
    )

    doc.add_paragraph()

    # ==================== TAB 1 ====================
    doc.add_heading('3.1. Tab 1: Thông tin cơ bản', 2)

    doc.add_paragraph(
        'Tab này thu thập các thông tin cơ bản về hệ thống CNTT. '
        'Đây là tab đầu tiên và bắt buộc phải hoàn thành trước khi điền các tab khác.'
    )

    # Add screenshot
    add_image_with_caption(doc, f"{SCREENSHOTS_DIR}/tab1-co-ban-filled.png",
                          "Hình 1: Giao diện Tab 1 - Thông tin cơ bản")

    doc.add_heading('Các trường cần điền:', 3)

    tab1_fields = [
        ('Tổ chức (*)', 'Chọn đơn vị sở hữu/quản lý hệ thống từ danh sách dropdown'),
        ('Mã hệ thống', 'Tự động sinh bởi hệ thống, không cần nhập'),
        ('Tên hệ thống (*)', 'Nhập tên đầy đủ của hệ thống bằng tiếng Việt'),
        ('Tên tiếng Anh (*)', 'Nhập tên hệ thống bằng tiếng Anh (nếu không có, nhập tên tiếng Việt)'),
        ('Mô tả (*)', 'Mô tả ngắn gọn chức năng và mục đích của hệ thống'),
        ('Trạng thái (*)', 'Chọn trạng thái hiện tại: Đang vận hành, Thí điểm, Bảo trì, v.v.'),
        ('Mức độ quan trọng (*)', 'Đánh giá mức độ quan trọng: Cực kỳ quan trọng, Quan trọng, Trung bình, Thấp'),
        ('Phạm vi sử dụng (*)', 'Chọn: Nội bộ (chỉ dùng trong đơn vị) hoặc Bên ngoài (công dân/doanh nghiệp sử dụng)'),
        ('Nhu cầu (*)', 'Chọn loại nhu cầu: Tích hợp - Liên thông, Nâng cấp, Thay thế, v.v.'),
        ('Thời gian mong muốn hoàn thành (*)', 'Chọn tháng/năm dự kiến hoàn thành'),
        ('Nhóm hệ thống (*)', 'Phân loại hệ thống: CSDL chuyên ngành, Ứng dụng nội bộ, v.v.'),
        ('Ghi chú bổ sung (*)', 'Thông tin bổ sung cho tab này'),
    ]

    for field, desc in tab1_fields:
        p = doc.add_paragraph()
        p.add_run(f'• {field}: ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ==================== TAB 2 ====================
    doc.add_heading('3.2. Tab 2: Bối cảnh nghiệp vụ', 2)

    doc.add_paragraph(
        'Tab này thu thập thông tin về mục tiêu nghiệp vụ, đối tượng sử dụng '
        'và các chỉ số về người dùng của hệ thống.'
    )

    add_image_with_caption(doc, f"{SCREENSHOTS_DIR}/tab2-nghiep-vu-filled.png",
                          "Hình 2: Giao diện Tab 2 - Bối cảnh nghiệp vụ")

    doc.add_heading('Các trường cần điền:', 3)

    tab2_fields = [
        ('Mục tiêu nghiệp vụ (*)', 'Chọn một hoặc nhiều mục tiêu: Số hóa quy trình, Cải thiện dịch vụ công, Tăng cường minh bạch, v.v.'),
        ('Quy trình nghiệp vụ chính (*)', 'Chọn các quy trình: Quản lý hồ sơ, Phê duyệt, Tra cứu, Báo cáo, Thống kê, v.v.'),
        ('Có đủ hồ sơ phân tích thiết kế? (*)', 'Bật/Tắt - cho biết hệ thống có tài liệu thiết kế đầy đủ không'),
        ('Đối tượng sử dụng (*)', 'Chọn các loại người dùng: Cán bộ nội bộ, Doanh nghiệp, Người dân, Địa phương, v.v.'),
        ('Số lượng người dùng hàng năm (*)', 'Nhập số người dùng ước tính trong năm'),
        ('Tổng số tài khoản (*)', 'Tổng số tài khoản đã đăng ký trên hệ thống'),
        ('MAU - Monthly Active Users (*)', 'Số người dùng hoạt động trong tháng'),
        ('DAU - Daily Active Users (*)', 'Số người dùng hoạt động trong ngày'),
        ('Số đơn vị/địa phương sử dụng (*)', 'Số lượng đơn vị hoặc địa phương đang sử dụng hệ thống'),
        ('Ghi chú bổ sung (*)', 'Thông tin bổ sung cho tab này'),
    ]

    for field, desc in tab2_fields:
        p = doc.add_paragraph()
        p.add_run(f'• {field}: ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ==================== TAB 3 ====================
    doc.add_heading('3.3. Tab 3: Kiến trúc công nghệ', 2)

    doc.add_paragraph(
        'Tab này thu thập thông tin chi tiết về công nghệ sử dụng trong hệ thống. '
        'Đây là tab có nhiều trường nhất và yêu cầu người nhập có kiến thức về CNTT.'
    )

    add_image_with_caption(doc, f"{SCREENSHOTS_DIR}/tab3-cong-nghe-filled.png",
                          "Hình 3: Giao diện Tab 3 - Kiến trúc công nghệ")

    doc.add_heading('Các trường cần điền:', 3)

    tab3_fields = [
        ('Ngôn ngữ lập trình (*)', 'Chọn một hoặc nhiều: Python, Java, JavaScript, C#, PHP, v.v.'),
        ('Framework/Thư viện (*)', 'Chọn frameworks: Django, Spring Boot, React, Angular, Vue.js, Laravel, v.v.'),
        ('Cơ sở dữ liệu (*)', 'Chọn CSDL đang sử dụng: SQL Server, MySQL, PostgreSQL, Oracle, MongoDB, v.v.'),
        ('Nền tảng triển khai (*)', 'Chọn: On-premise, AWS, Azure, Google Cloud, v.v.'),
        ('Backend Technology (*)', 'Công nghệ phía máy chủ: Node.js, Python, Java, C#/.NET, v.v.'),
        ('Frontend Technology (*)', 'Công nghệ phía giao diện: React, Vue.js, Angular, HTML/CSS/JS thuần, v.v.'),
        ('Loại kiến trúc (*)', 'Chọn: Monolithic, Microservices, Serverless, SaaS, v.v.'),
        ('Container hóa (*)', 'Chọn: Docker, Kubernetes, OpenShift, hoặc "Không sử dụng"'),
        ('Multi-tenant (*)', 'Hệ thống có hỗ trợ nhiều đơn vị thuê chung không?'),
        ('Phân lớp (Layered) (*)', 'Hệ thống có kiến trúc phân lớp không?'),
        ('Chi tiết phân lớp (*)', 'Mô tả các lớp: Presentation, Business Logic, Data Access, v.v.'),
        ('API Style (*)', 'Chọn: REST API, GraphQL, gRPC, SOAP, hoặc ghi rõ nếu khác'),
        ('Messaging/Queue (*)', 'Chọn: Apache Kafka, RabbitMQ, ActiveMQ, hoặc "Không sử dụng"'),
        ('Cache System (*)', 'Chọn: Redis, Memcached, hoặc "Không sử dụng"'),
        ('Search Engine (*)', 'Chọn: Elasticsearch, Solr, hoặc "Không sử dụng"'),
        ('Reporting/BI Tool (*)', 'Chọn: Power BI, Tableau, Tự phát triển, v.v.'),
        ('Source Repository (*)', 'Chọn: GitHub, GitLab, Bitbucket, hoặc "Không quản lý"'),
        ('CI/CD Pipeline (*)', 'Hệ thống có quy trình CI/CD không?'),
        ('CI/CD Tool (*)', 'Chọn: Jenkins, GitHub Actions, GitLab CI, v.v.'),
        ('Automated Testing (*)', 'Hệ thống có kiểm thử tự động không?'),
        ('Testing Tools (*)', 'Nhập công cụ: Jest, Pytest, Selenium, JUnit, v.v.'),
        ('Ghi chú bổ sung (*)', 'Thông tin bổ sung cho tab này'),
    ]

    for field, desc in tab3_fields:
        p = doc.add_paragraph()
        p.add_run(f'• {field}: ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ==================== TAB 4 ====================
    doc.add_heading('3.4. Tab 4: Kiến trúc dữ liệu', 2)

    doc.add_paragraph(
        'Tab này thu thập thông tin về cách quản lý và lưu trữ dữ liệu trong hệ thống.'
    )

    add_image_with_caption(doc, f"{SCREENSHOTS_DIR}/tab4-du-lieu-filled.png",
                          "Hình 4: Giao diện Tab 4 - Kiến trúc dữ liệu")

    doc.add_heading('Các trường cần điền:', 3)

    tab4_fields = [
        ('Nguồn dữ liệu (*)', 'Chọn: User Input, External APIs, Database Sync, File Import, v.v.'),
        ('Loại dữ liệu (*)', 'Chọn: Business Data, Documents, Statistics, Master Data, v.v.'),
        ('Phân loại dữ liệu (*)', 'Mức độ bảo mật: Công khai, Nội bộ, Hạn chế, Bí mật, Tối mật'),
        ('Khối lượng dữ liệu (*)', 'Ước tính tổng dung lượng: <1GB, 1-10GB, 10-100GB, 100GB-1TB, >1TB'),
        ('Dung lượng CSDL hiện tại (GB) (*)', 'Nhập số GB dung lượng database'),
        ('Dung lượng file đính kèm (GB) (*)', 'Nhập số GB dung lượng file'),
        ('Tốc độ tăng trưởng dữ liệu (%) (*)', 'Tốc độ tăng trưởng dữ liệu hàng năm'),
        ('Loại lưu trữ file (*)', 'Chọn: File Server, Object Storage, NAS, Database BLOB, v.v.'),
        ('Số bản ghi (*)', 'Tổng số bản ghi trong CSDL'),
        ('CSDL phụ/khác (*)', 'Các CSDL phụ hoặc cache: MySQL, Redis, MongoDB, v.v.'),
        ('Chính sách lưu trữ dữ liệu', 'Mô tả chính sách: thời gian lưu, archive, xóa dữ liệu'),
        ('Data Catalog (*)', 'Có danh mục dữ liệu không?'),
        ('Master Data Management (*)', 'Có hệ thống quản lý dữ liệu gốc không?'),
    ]

    for field, desc in tab4_fields:
        p = doc.add_paragraph()
        p.add_run(f'• {field}: ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ==================== TAB 5 ====================
    doc.add_heading('3.5. Tab 5: Tích hợp hệ thống', 2)

    doc.add_paragraph(
        'Tab này thu thập thông tin về cách hệ thống tích hợp với các hệ thống khác '
        'trong và ngoài đơn vị.'
    )

    add_image_with_caption(doc, f"{SCREENSHOTS_DIR}/tab5-tich-hop-filled.png",
                          "Hình 5: Giao diện Tab 5 - Tích hợp hệ thống")

    doc.add_heading('Các trường cần điền:', 3)

    tab5_fields = [
        ('Số API cung cấp (*)', 'Số lượng API mà hệ thống cung cấp cho bên ngoài'),
        ('Số API tiêu thụ (*)', 'Số lượng API mà hệ thống gọi từ bên ngoài'),
        ('Chuẩn API (*)', 'Chọn: OpenAPI 3.0, OpenAPI 2.0, SOAP WSDL, GraphQL, gRPC, v.v.'),
        ('Có API Gateway? (*)', 'Hệ thống có sử dụng API Gateway không?'),
        ('Tên API Gateway (*)', 'Chọn: Kong, AWS API Gateway, Azure APIM, Không có, v.v.'),
        ('Có API Versioning? (*)', 'Có quản lý phiên bản API không?'),
        ('Có Rate Limiting? (*)', 'Có giới hạn số lượng request không?'),
        ('Tài liệu API (*)', 'Link hoặc mô tả tài liệu API (Swagger/OpenAPI docs)'),
        ('Chuẩn phiên bản API (*)', 'Chọn: URL Path, Header, Query Parameter, v.v.'),
        ('Có giám sát tích hợp? (*)', 'Có hệ thống giám sát các kết nối tích hợp không?'),
        ('Hệ thống nội bộ tích hợp (*)', 'Chọn: Quản lý văn bản, Nhân sự, Tài chính, Portal nội bộ, v.v.'),
        ('Hệ thống bên ngoài tích hợp (*)', 'Chọn: VNeID, LGSP, Cổng DVC Quốc gia, Thuế, Hải quan, v.v.'),
        ('API/Webservices (*)', 'Liệt kê các endpoint API/webservices'),
        ('Phương thức trao đổi dữ liệu (*)', 'Chọn: API REST, API SOAP, File Transfer, Database Link, v.v.'),
    ]

    for field, desc in tab5_fields:
        p = doc.add_paragraph()
        p.add_run(f'• {field}: ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ==================== TAB 6 ====================
    doc.add_heading('3.6. Tab 6: An toàn thông tin', 2)

    doc.add_paragraph(
        'Tab này thu thập thông tin về các biện pháp bảo mật và an toàn thông tin của hệ thống.'
    )

    doc.add_heading('Các trường cần điền:', 3)

    tab6_fields = [
        ('Phương thức xác thực (*)', 'Chọn: Username/Password, SSO, 2FA, OAuth, LDAP, v.v.'),
        ('Có mã hóa dữ liệu? (*)', 'Hệ thống có mã hóa dữ liệu nhạy cảm không?'),
        ('Có log audit trail? (*)', 'Hệ thống có ghi lại nhật ký hoạt động không?'),
        ('Mức độ an toàn thông tin (*)', 'Chọn cấp độ: Cấp 1 đến Cấp 5 theo quy định'),
        ('Có tài liệu ATTT? (*)', 'Hệ thống có đầy đủ tài liệu về an toàn thông tin không?'),
        ('Ghi chú bổ sung (*)', 'Thông tin bổ sung về bảo mật'),
    ]

    for field, desc in tab6_fields:
        p = doc.add_paragraph()
        p.add_run(f'• {field}: ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ==================== TAB 7 ====================
    doc.add_heading('3.7. Tab 7: Hạ tầng', 2)

    doc.add_paragraph(
        'Tab này thu thập thông tin về hạ tầng kỹ thuật phục vụ vận hành hệ thống.'
    )

    doc.add_heading('Các trường cần điền:', 3)

    tab7_fields = [
        ('Cấu hình server (*)', 'Mô tả cấu hình máy chủ: CPU, RAM, Storage'),
        ('Phương án backup (*)', 'Chọn: Daily backup, Weekly backup, Real-time replication, v.v.'),
        ('Dung lượng lưu trữ (*)', 'Tổng dung lượng storage được cấp'),
        ('Kế hoạch phục hồi thảm họa (*)', 'Mô tả kế hoạch DR (Disaster Recovery)'),
        ('Vị trí triển khai (*)', 'Chọn: Data Center Bộ, Data Center thuê, Cloud, v.v.'),
        ('Cấu hình tính toán (*)', 'Chi tiết về VM/Container/Physical server'),
        ('Loại hạ tầng tính toán (*)', 'Chọn: VM, Container, Physical, Serverless'),
        ('Tần suất triển khai (*)', 'Chọn: Hàng ngày, Hàng tuần, Hàng tháng, Khi cần'),
        ('Ghi chú bổ sung (*)', 'Thông tin bổ sung về hạ tầng'),
    ]

    for field, desc in tab7_fields:
        p = doc.add_paragraph()
        p.add_run(f'• {field}: ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ==================== TAB 8 ====================
    doc.add_heading('3.8. Tab 8: Vận hành', 2)

    doc.add_paragraph(
        'Tab này thu thập thông tin về người phụ trách và cách thức vận hành hệ thống.'
    )

    doc.add_heading('Các trường cần điền:', 3)

    tab8_fields = [
        ('Người chịu trách nhiệm nghiệp vụ (*)', 'Họ tên người phụ trách về mặt nghiệp vụ'),
        ('Người chịu trách nhiệm kỹ thuật (*)', 'Họ tên người phụ trách về mặt kỹ thuật'),
        ('Người chịu trách nhiệm (*)', 'Họ tên người liên hệ chính'),
        ('Số điện thoại liên hệ (*)', 'Số điện thoại để liên hệ khi cần hỗ trợ'),
        ('Email liên hệ (*)', 'Email để liên hệ'),
        ('Mức độ hỗ trợ (*)', 'Chọn: 24/7, Giờ hành chính, Theo yêu cầu, v.v.'),
        ('Ghi chú bổ sung (*)', 'Thông tin bổ sung về vận hành'),
    ]

    for field, desc in tab8_fields:
        p = doc.add_paragraph()
        p.add_run(f'• {field}: ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ==================== TAB 9 ====================
    doc.add_heading('3.9. Tab 9: Đánh giá', 2)

    doc.add_paragraph(
        'Tab này thu thập các đánh giá về hiệu năng, chất lượng và kế hoạch tương lai của hệ thống.'
    )

    doc.add_heading('Các trường cần điền:', 3)

    tab9_fields = [
        ('Đánh giá hiệu năng (*)', 'Chọn: Rất tốt, Tốt, Trung bình, Kém, Rất kém'),
        ('Đánh giá hài lòng người dùng (*)', 'Chọn: Rất hài lòng, Hài lòng, Bình thường, Không hài lòng'),
        ('Mức độ nợ kỹ thuật (*)', 'Chọn: Thấp, Trung bình, Cao, Rất cao'),
        ('Đề xuất hành động (*)', 'Chọn: Duy trì, Nâng cấp, Thay thế, Loại bỏ'),
        ('Điểm phù hợp cho tích hợp (*)', 'Chọn các điểm mạnh của hệ thống về tích hợp'),
        ('Điểm vướng mắc (*)', 'Chọn các vấn đề gặp phải'),
        ('Uptime (%) (*)', 'Tỷ lệ thời gian hoạt động (ví dụ: 99.9)'),
        ('Thời gian phản hồi trung bình (ms) (*)', 'Thời gian response trung bình'),
        ('Kế hoạch thay thế (*)', 'Mô tả kế hoạch thay thế nếu có (nếu không có nhập "Không có")'),
        ('Các vấn đề chính (*)', 'Liệt kê các vấn đề gặp phải (nếu không có nhập "Không có")'),
        ('Đề xuất cải tiến (*)', 'Các đề xuất để cải tiến hệ thống'),
        ('Kế hoạch tương lai (*)', 'Mô tả kế hoạch phát triển'),
        ('Mức độ ưu tiên hiện đại hóa (*)', 'Chọn: Cao, Trung bình, Thấp'),
    ]

    for field, desc in tab9_fields:
        p = doc.add_paragraph()
        p.add_run(f'• {field}: ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ==================== 4. LƯU Ý QUAN TRỌNG ====================
    doc.add_heading('4. LƯU Ý QUAN TRỌNG', 1)

    important_notes = [
        ('Điền đầy đủ tất cả các trường bắt buộc',
         'Mỗi tab đều có các trường bắt buộc (đánh dấu *). Hệ thống sẽ KHÔNG cho phép '
         'chuyển tab hoặc lưu dữ liệu nếu còn trường bắt buộc chưa điền.'),

        ('Quy trình nhập liệu theo thứ tự',
         'Phải hoàn thành từng tab theo thứ tự từ 1 đến 9. Không thể bỏ qua tab nào.'),

        ('Lưu thường xuyên',
         'Sau khi hoàn thành mỗi tab, nhấn nút "Lưu & Tiếp tục" để lưu dữ liệu và chuyển '
         'sang tab tiếp theo. Điều này giúp tránh mất dữ liệu.'),

        ('Không biết thông tin',
         'Nếu không biết chính xác thông tin, hãy nhập giá trị ước tính hoặc "Không có"/"Không rõ" '
         'tùy theo loại trường. Không được để trống trường bắt buộc.'),

        ('Hỏi hỗ trợ kỹ thuật',
         'Với các trường về công nghệ (Tab 3-5), nếu không chắc chắn, hãy liên hệ bộ phận '
         'kỹ thuật của đơn vị để có thông tin chính xác.'),
    ]

    for title, desc in important_notes:
        p = doc.add_paragraph()
        p.add_run(f'✓ {title}: ').bold = True
        p.add_run(desc)
        doc.add_paragraph()

    doc.add_page_break()

    # ==================== 5. XỬ LÝ LỖI ====================
    doc.add_heading('5. XỬ LÝ LỖI THƯỜNG GẶP', 1)

    errors = [
        ('Thông báo "Vui lòng điền đủ X trường bắt buộc..."',
         'Nguyên nhân: Còn trường bắt buộc chưa điền trong tab hiện tại.\n'
         'Cách xử lý: Kiểm tra các trường có viền đỏ và thông báo lỗi, điền đầy đủ thông tin.'),

        ('Không thể chuyển sang tab tiếp theo',
         'Nguyên nhân: Tab hiện tại chưa hoàn thành.\n'
         'Cách xử lý: Điền đầy đủ tất cả các trường bắt buộc trong tab hiện tại, sau đó nhấn '
         '"Lưu & Tiếp tục".'),

        ('Nút "Lưu" bị mờ (disabled)',
         'Nguyên nhân: Dữ liệu chưa thay đổi hoặc form chưa hợp lệ.\n'
         'Cách xử lý: Kiểm tra và điền đầy đủ các trường bắt buộc.'),

        ('Phiên làm việc hết hạn',
         'Nguyên nhân: Không hoạt động trong thời gian dài.\n'
         'Cách xử lý: Đăng nhập lại. Có thể tích "Ghi nhớ đăng nhập" để duy trì phiên.'),

        ('Dữ liệu không lưu được',
         'Nguyên nhân: Lỗi kết nối hoặc server.\n'
         'Cách xử lý: Thử lại sau vài phút. Nếu vẫn lỗi, liên hệ quản trị viên.'),
    ]

    for error_title, error_desc in errors:
        doc.add_heading(error_title, 2)
        for line in error_desc.split('\n'):
            doc.add_paragraph(line)

    # ==================== FOOTER ====================
    doc.add_page_break()
    doc.add_heading('LIÊN HỆ HỖ TRỢ', 1)

    doc.add_paragraph('Trung tâm Công nghệ thông tin - Bộ Khoa học và Công nghệ')
    doc.add_paragraph('Địa chỉ: 113 Trần Duy Hưng, Cầu Giấy, Hà Nội')
    doc.add_paragraph('Email: support@mst.gov.vn')
    doc.add_paragraph('Hotline: 024-xxx-xxxx')

    doc.add_paragraph()
    doc.add_paragraph()

    footer_note = doc.add_paragraph('Tài liệu được cập nhật: Tháng 01/2026')
    footer_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_note.runs[0].italic = True

    # Save document
    output_path = "Huong_Dan_Su_Dung_He_Thong_Khao_Sat_CDS.docx"
    doc.save(output_path)
    print(f"✓ Tài liệu đã được tạo: {output_path}")
    return output_path

if __name__ == "__main__":
    create_user_guide()
